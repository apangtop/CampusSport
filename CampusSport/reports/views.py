import io
from rest_framework import viewsets, status, permissions
from rest_framework.decorators import action
from rest_framework.response import Response
from django.http import HttpResponse

# Word
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Excel
import openpyxl
from openpyxl.styles import Font as XFont, Alignment as XAlignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

from events.models import SportsMeet, Event, Schedule
from registration.models import Registration
from scores.models import Score, ClassPoints
from .models import Report


class IsAdmin(permissions.BasePermission):
    def has_permission(self, request, view):
        return request.user.is_authenticated and request.user.role == 'admin'


# ─────────────────────────────────────────────
#  Word 工具函数
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    """设置单元格四边细线"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_cell_text(cell, text, bold=False, font_size=10, color=None,
                  align=WD_ALIGN_PARAGRAPH.CENTER, v_align=WD_ALIGN_VERTICAL.CENTER):
    """向单元格写入文字并设置样式"""
    cell.vertical_alignment = v_align
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text) if text is not None else '')
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_centered_para(doc, text, font_size=12, bold=False, space_before=0, space_after=6, color=None):
    """添加居中段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_heading_para(doc, text, font_size=13, bold=True, left_bar=True):
    """添加带色块的标题段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    if left_bar:
        # 左侧色条用 pPr border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '1a6db5')
        pBdr.append(left)
        pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(26, 109, 181)
    return p


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def set_table_style(table):
    """全表边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'AAAAAA')
        tblBorders.append(b)
    tblPr.append(tblBorders)


# ─────────────────────────────────────────────
#  秩序册 Word 生成
# ─────────────────────────────────────────────

def build_order_book_word(sports_meet):
    doc = Document()

    # 页面设置：A4，页边距2cm
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ══════════════════════════════════════
    #  封面
    # ══════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    add_centered_para(doc, sports_meet.school or '学校名称', font_size=16, space_after=12)
    add_centered_para(doc, f'第{sports_meet.session}届 {sports_meet.name}',
                      font_size=26, bold=True, space_after=16)

    title_p = add_centered_para(doc, '秩  序  册', font_size=40, bold=True, space_after=24)
    title_p.runs[0].font.color.rgb = RGBColor(26, 109, 181)

    if sports_meet.start_date:
        date_str = sports_meet.start_date.strftime('%Y年%m月%d日')
        if sports_meet.end_date and sports_meet.end_date != sports_meet.start_date:
            date_str += f' — {sports_meet.end_date.strftime("%m月%d日")}'
        add_centered_para(doc, f'举办日期：{date_str}', font_size=14, space_before=8)

    # 分页
    doc.add_page_break()

    # ══════════════════════════════════════
    #  目录
    # ══════════════════════════════════════
    add_heading_para(doc, '目  录', font_size=16, left_bar=False)
    doc.add_paragraph()

    events = sports_meet.events.prefetch_related('schedules').select_related('referee').all()
    for idx, event in enumerate(events, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'{idx:02d}.  {event.name}（{event.get_gender_display()}）')
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════
    #  赛程总览表
    # ══════════════════════════════════════
    add_heading_para(doc, '赛程总览', font_size=15)
    doc.add_paragraph()

    overview_headers = ['项目名称', '性别', '阶段', '组次', '比赛时间', '场地', '负责裁判']
    overview_widths = [4.0, 1.4, 1.8, 1.4, 3.5, 3.5, 2.4]
    header_bg = '1a6db5'

    overview_rows = []
    for event in events:
        scheds = list(event.schedules.all())
        referee_name = event.referee.real_name if event.referee else '待定'
        if not scheds:
            overview_rows.append([
                event.name, event.get_gender_display(), '直接决赛', '-', '-', '-', referee_name
            ])
        else:
            for sch in scheds:
                time_str = sch.scheduled_time.strftime('%Y-%m-%d %H:%M') if sch.scheduled_time else '-'
                overview_rows.append([
                    event.name, event.get_gender_display(),
                    sch.get_stage_display(), f'第{sch.group_number}组',
                    time_str, sch.venue or '-', referee_name
                ])

    table = doc.add_table(rows=1 + len(overview_rows), cols=len(overview_headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_style(table)

    # 表头
    for col_idx, (h, w) in enumerate(zip(overview_headers, overview_widths)):
        cell = table.cell(0, col_idx)
        cell.width = Cm(w)
        set_cell_bg(cell, header_bg)
        add_cell_text(cell, h, bold=True, font_size=10, color=(255, 255, 255))

    # 数据行
    for r_idx, row_data in enumerate(overview_rows, 1):
        bg = 'F5F7FA' if r_idx % 2 == 0 else 'FFFFFF'
        for col_idx, val in enumerate(row_data):
            cell = table.cell(r_idx, col_idx)
            cell.width = Cm(overview_widths[col_idx])
            set_cell_bg(cell, bg)
            add_cell_text(cell, val, font_size=9.5)

    doc.add_page_break()

    # ══════════════════════════════════════
    #  各项目参赛名单
    # ══════════════════════════════════════
    for event in events:
        # 项目标题
        add_heading_para(doc, f'▶  {event.name}', font_size=14)

        # 项目信息行
        referee_name = event.referee.real_name if event.referee else '待定'
        # 汇总赛程时间、场地
        scheds = event.schedules.all()
        times = [s.scheduled_time.strftime('%m/%d %H:%M') for s in scheds if s.scheduled_time]
        venues = list(set(s.venue for s in scheds if s.venue))
        time_str = '、'.join(times) if times else '待定'
        venue_str = '、'.join(venues) if venues else '待定'

        info_parts = [
            f'性别：{event.get_gender_display()}',
            f'成绩单位：{event.get_result_unit_display()}',
            f'赛制：{event.get_stage_type_display()}',
            f'比赛时间：{time_str}',
            f'场地：{venue_str}',
            f'负责裁判：{referee_name}',
        ]
        info_p = doc.add_paragraph('    '.join(info_parts))
        info_p.paragraph_format.space_after = Pt(6)
        for run in info_p.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)

        # 获取报名数据，按赛程分组
        registrations = Registration.objects.filter(
            event=event, status__in=['submitted', 'approved']
        ).select_related('student', 'schedule').order_by('schedule__group_number', 'lane', 'student__class_name')

        # 按 schedule 分组
        groups = {}
        no_sch_regs = []
        for reg in registrations:
            if reg.schedule:
                key = reg.schedule.id
                if key not in groups:
                    groups[key] = {'schedule': reg.schedule, 'regs': []}
                groups[key]['regs'].append(reg)
            else:
                no_sch_regs.append(reg)

        all_groups = list(groups.values())
        if no_sch_regs:
            all_groups.append({'schedule': None, 'regs': no_sch_regs})

        if not all_groups:
            p = doc.add_paragraph('  暂无报名人员')
            p.paragraph_format.space_after = Pt(8)
        else:
            for grp in all_groups:
                sch = grp['schedule']
                grp_regs = grp['regs']

                # 组次信息子标题
                if sch:
                    time_str = sch.scheduled_time.strftime('%Y-%m-%d %H:%M') if sch.scheduled_time else '时间待定'
                    sub_title = f'【{sch.get_stage_display()} 第{sch.group_number}组】  {time_str}  {sch.venue or ""}'
                else:
                    sub_title = '【参赛名单】'

                sp = doc.add_paragraph(sub_title)
                sp.paragraph_format.space_before = Pt(4)
                sp.paragraph_format.space_after = Pt(3)
                for run in sp.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(10)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(44, 62, 80)

                # 参赛名单表
                reg_headers = ['道次', '姓名', '班级', '成绩', '名次']
                reg_widths = [2.0, 2.8, 4.0, 3.0, 2.0]

                reg_table = doc.add_table(rows=1 + len(grp_regs), cols=len(reg_headers))
                reg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_style(reg_table)

                for col_idx, (h, w) in enumerate(zip(reg_headers, reg_widths)):
                    cell = reg_table.cell(0, col_idx)
                    cell.width = Cm(w)
                    set_cell_bg(cell, '2C3E50')
                    add_cell_text(cell, h, bold=True, font_size=10, color=(255, 255, 255))

                for r_idx, reg in enumerate(grp_regs, 1):
                    score_obj = Score.objects.filter(registration=reg, stage='final').first()
                    result_val = score_obj.result if score_obj else ''
                    rank_val = f'第{score_obj.rank}名' if (score_obj and score_obj.rank) else ''
                    bg = 'F8F9FA' if r_idx % 2 == 0 else 'FFFFFF'
                    row_data = [reg.lane or '', reg.student.name,
                                reg.student.class_name, result_val, rank_val]
                    for col_idx, val in enumerate(row_data):
                        cell = reg_table.cell(r_idx, col_idx)
                        cell.width = Cm(reg_widths[col_idx])
                        set_cell_bg(cell, bg)
                        add_cell_text(cell, val, font_size=10)

                doc.add_paragraph()

        doc.add_page_break()

    # ══════════════════════════════════════
    #  积分规则
    # ══════════════════════════════════════
    add_heading_para(doc, '积分规则说明', font_size=15)
    doc.add_paragraph()

    rules_headers = ['名次', '积分']
    rules_widths = [4.0, 4.0]
    rules_data = [(1, 7), (2, 5), (3, 4), (4, 3), (5, 2), (6, 1), ('第7名及以后', 0)]

    rules_table = doc.add_table(rows=1 + len(rules_data), cols=2)
    rules_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_style(rules_table)

    for col_idx, h in enumerate(rules_headers):
        cell = rules_table.cell(0, col_idx)
        cell.width = Cm(rules_widths[col_idx])
        set_cell_bg(cell, '1a6db5')
        add_cell_text(cell, h, bold=True, font_size=11, color=(255, 255, 255))

    for r_idx, (rank, pts) in enumerate(rules_data, 1):
        bg = 'F5F7FA' if r_idx % 2 == 0 else 'FFFFFF'
        rank_str = f'第{rank}名' if isinstance(rank, int) else rank
        for col_idx, val in enumerate([rank_str, f'{pts} 分']):
            cell = rules_table.cell(r_idx, col_idx)
            cell.width = Cm(rules_widths[col_idx])
            set_cell_bg(cell, bg)
            add_cell_text(cell, val, font_size=11, bold=(r_idx <= 3))

    return doc


# ─────────────────────────────────────────────
#  成绩报表 Excel（保持不变）
# ─────────────────────────────────────────────

def build_result_report_excel(sports_meet):
    wb = openpyxl.Workbook()
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
    header_font_white = XFont(name='微软雅黑', size=11, bold=True, color='FFFFFF')
    normal_font = XFont(name='微软雅黑', size=10)
    center_align = XAlignment(horizontal='center', vertical='center')

    # 班级积分榜
    points_ws = wb.active
    points_ws.title = '班级积分榜'
    headers = ['排名', '班级', '总积分', '金牌', '银牌', '铜牌']
    widths = [8, 20, 12, 8, 8, 8]
    for col, (h, w) in enumerate(zip(headers, widths), 1):
        c = points_ws.cell(row=1, column=col, value=h)
        c.font = header_font_white
        c.fill = header_fill
        c.alignment = center_align
        c.border = thin_border
        points_ws.column_dimensions[get_column_letter(col)].width = w

    class_points = ClassPoints.objects.filter(sports_meet=sports_meet).order_by('rank')
    for r_idx, cp in enumerate(class_points, 2):
        for col, val in enumerate([cp.rank, cp.class_name, cp.total_points,
                                   cp.gold_medals, cp.silver_medals, cp.bronze_medals], 1):
            c = points_ws.cell(row=r_idx, column=col, value=val)
            c.font = normal_font
            c.alignment = center_align
            c.border = thin_border

    # 各项目成绩页
    events = sports_meet.events.select_related('referee').all()
    for event in events:
        ws = wb.create_sheet(event.name[:28])
        ws.merge_cells('A1:F1')
        ws['A1'] = f'{event.name} 成绩单'
        ws['A1'].font = XFont(name='微软雅黑', size=13, bold=True)
        ws['A1'].alignment = center_align

        hdrs = ['名次', '姓名', '班级', '成绩', '积分', '阶段']
        wids = [8, 15, 20, 15, 10, 10]
        for col, (h, w) in enumerate(zip(hdrs, wids), 1):
            c = ws.cell(row=2, column=col, value=h)
            c.font = header_font_white
            c.fill = header_fill
            c.alignment = center_align
            c.border = thin_border
            ws.column_dimensions[get_column_letter(col)].width = w

        scores = Score.objects.filter(
            registration__event=event
        ).select_related('registration__student').order_by('stage', 'rank')
        for r_idx, score in enumerate(scores, 3):
            for col, val in enumerate([
                score.rank or '-', score.registration.student.name,
                score.registration.student.class_name, score.result or '-',
                score.points, score.get_stage_display()
            ], 1):
                c = ws.cell(row=r_idx, column=col, value=val)
                c.font = normal_font
                c.alignment = center_align
                c.border = thin_border

    return wb


# ─────────────────────────────────────────────
#  ViewSet
# ─────────────────────────────────────────────

class ReportViewSet(viewsets.ModelViewSet):
    queryset = Report.objects.all()
    permission_classes = [IsAdmin]

    def get_queryset(self):
        qs = super().get_queryset()
        meet_id = self.request.query_params.get('sports_meet')
        if meet_id:
            qs = qs.filter(sports_meet_id=meet_id)
        return qs

    @action(detail=False, methods=['get'])
    def generate_order_book(self, request):
        """生成并下载秩序册 Word (.docx)"""
        meet_id = request.query_params.get('sports_meet_id')
        if not meet_id:
            return Response({'detail': '需要 sports_meet_id 参数'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            meet = SportsMeet.objects.get(id=meet_id)
        except SportsMeet.DoesNotExist:
            return Response({'detail': '运动会不存在'}, status=status.HTTP_404_NOT_FOUND)

        doc = build_order_book_word(meet)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f'秩序册_{meet.name}_第{meet.session}届.docx'
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = (
            f'attachment; filename="{filename.encode("utf-8").decode("latin-1")}"'
        )
        response['Access-Control-Expose-Headers'] = 'Content-Disposition'
        return response

    @action(detail=False, methods=['get'])
    def generate_result_report(self, request):
        """生成并下载成绩报表 Excel"""
        meet_id = request.query_params.get('sports_meet_id')
        if not meet_id:
            return Response({'detail': '需要 sports_meet_id 参数'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            meet = SportsMeet.objects.get(id=meet_id)
        except SportsMeet.DoesNotExist:
            return Response({'detail': '运动会不存在'}, status=status.HTTP_404_NOT_FOUND)

        wb = build_result_report_excel(meet)
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        filename = f'成绩报表_{meet.name}_第{meet.session}届.xlsx'
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = (
            f'attachment; filename="{filename.encode("utf-8").decode("latin-1")}"'
        )
        response['Access-Control-Expose-Headers'] = 'Content-Disposition'
        return response
